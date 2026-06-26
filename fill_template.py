#!/usr/bin/env python3
"""
Fill 聂琦钢's dev log from PDF into the Word template format.
"""

from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

TEMPLATE = "个人开发日志模板.docx"
OUTPUT = "聂琦钢_个人开发日志.docx"
IMG_DIR = "pdf_images"

DAILY_DATA = {
    "2026.06.18": {
        "工作计划": "1. 初始化项目工程环境（package.json / tsconfig / vite / tailwind / postcss）\n"
                     "2. 设计核心类型系统（types/algo.ts 第一版）\n"
                     "3. 实现全局状态机（AlgoContext.tsx / Zustand Store）\n"
                     "4. 编写 Electron 桌面壳（main.ts）\n"
                     "5. 编写入口文件与全局样式（main.tsx / App.tsx / index.css）",
        "完成情况": "1. 初始化项目：npm init，配置 package.json，安装 React 18 / TypeScript 5 / Vite 5 / Tailwind CSS 3 / Zustand 4 / Framer Motion 11\n"
                     "2. 配置 tsconfig.json（strict 模式、路径别名 @/）、tsconfig.node.json（Electron 主进程编译）、vite.config.ts（路径解析、端口 5173）\n"
                     "3. 配置 tailwind.config.mjs（IDE 暗色主题色板：ide-bg/ide-sidebar/ide-panel/ide-border/ide-text/canvas-bg，自定义 keyframes：pulseHighlight/swapJump/fadeIn）\n"
                     "4. 配置 postcss.config.mjs（tailwindcss + autoprefixer）\n"
                     "5. 编写 index.html（JetBrains Mono + Inter 字体 CDN、中文 lang、暗色 body）\n"
                     "6. 编写 src/index.css（Tailwind 指令 + 代码高亮 token 颜色类 + 滚动条样式 + reduced-motion 适配）\n"
                     "7. 编写 src/types/algo.ts 第一版：AlgorithmCategory 联合类型、ActionType（6种）、VisualizerType（4种）、VisualStep 接口等\n"
                     "8. 编写 src/context/AlgoContext.tsx 第一版：14 个 Action 方法\n"
                     "9. 编写 electron/main.ts：BrowserWindow 创建、M1 macOS 适配、安全策略\n"
                     "10. 编写 src/main.tsx 和 src/App.tsx 入口文件",
        "出现的问题以及解决方案": "1. postcss.config.js 报 MODULE_TYPELESS_PACKAGE_JSON 警告 → 改为 .mjs 扩展名\n"
                               "2. npm install electron 二进制下载失败（socket hang up）→ 设置 ELECTRON_MIRROR 镜像\n"
                               "3. tsconfig.json 的 references 缺少 composite: true → 添加后与 declaration: false 冲突，最终删除 composite",
        "AI使用记录": "使用 Claude Code 进行项目脚手架搭建、类型接口生成、Zustand Store 实现\n"
                     "AI 生成代码约 800 行，人工修改约 20%（配置文件路径、样式细节、安全策略参数）\n"
                     "使用 GitHub Copilot 辅助 TypeScript 类型定义的补全",
    },
    "2026.06.19": {
        "工作计划": "1. 扩展类型系统第二版（ActionType 6→14种，VisualizerType 4→8种）\n"
                     "2. 更新 Canvas.tsx 和 Sidebar.tsx 适配新类型\n"
                     "3. 修复 generateRandomData 按算法 ID 定制",
        "完成情况": "1. 扩展 types/algo.ts 第二版：ActionType 从 6→14 种（新增 delete/recurse-in/recurse-out/divide/merge/relax/visit/backtrack）\n"
                     "2. VisualizerType 从 4→8 种（新增 array-dual/grid/grid+table/timeline）\n"
                     "3. VisualStep 新增扩展字段：secondaryData/graphEdges/gridData/gridCells/dpCellUpdates/pointerPositions/pointers\n"
                     "4. 新增辅助接口：GraphEdgeHighlight/GridCell/DPCellUpdate/PointerInfo\n"
                     "5. SpeedMultiplier 扩展到 5 档：0.25 | 0.5 | 1 | 2 | 4\n"
                     "6. 更新 Canvas.tsx：switch-case 分发新增 grid / grid+table 类型\n"
                     "7. 更新 Sidebar.tsx：categories 数组新增 'advanced'\n"
                     "8. 修复 generateRandomData 逻辑：从简单「全部随机数组」改为按算法 ID 定制",
        "出现的问题以及解决方案": "1. ActionType 扩展后，LogPanel 和 ArrayVisualizer 中 ACTION_LABELS/ACTION_COLORS 报类型错误 → 补全全部 16 种的中文标签和颜色\n"
                               "2. GraphEdgeHighlight 的 weight 字段类型不一致 → 统一使用可选 weight 字段\n"
                               "3. 部分算法（N皇后/迷宫）随机生成的数据无意义 → 按算法 ID 定制随机规则",
        "AI使用记录": "使用 Claude Code 进行类型系统扩展、ActionType 枚举扩展\n"
                     "AI 提供扩展方案，人工审核类型完备性",
    },
    "2026.06.20": {
        "工作计划": "1. 重写 randomData 生成器（按 55 个算法 ID 精确匹配）\n"
                     "2. 实现字符直输功能\n"
                     "3. 图算法边列表输入格式支持",
        "完成情况": "1. 重写 generateRandomData 函数：按 55 个算法 ID 精确匹配生成规则\n"
                     "   · 排序算法 → 8 个随机数\n"
                     "   · 二分查找/插值查找 → 10 个有序数 + 已有目标值\n"
                     "   · BST 查找 → 7 个 BST 节点 + 已有目标值\n"
                     "   · 顺表插入 → 5 元素 + 随机插入值 + 合法插入位\n"
                     "   · 顺表删除/链表删除 → 5 元素 + 合法删除位\n"
                     "   · 栈/队列 → 3-6 个随机值\n"
                     "   · N皇后 → 随机棋盘 4-8\n"
                     "   · 0/1背包 → 4 件随机重量 + 价值 + 合理容量\n"
                     "   · LCS → 2 组随机序列\n"
                     "   · 图算法 → 随机 4-7 节点 + 随机边 + 随机权重\n"
                     "   · 括号匹配 → 合法括号对再随机打乱\n"
                     "   · 中缀转后缀 → 合法数字表达式\n"
                     "2. 实现字符直输功能：DataControl 检测非数字字符时自动转 charCode\n"
                     "3. 图算法边列表输入：格式 N, from, to, w, ...\n"
                     "4. 更新 DataControl 的 isReadonly 逻辑和输入提示",
        "出现的问题以及解决方案": "1. 字符直输与数字输入模式切换逻辑需精确判断 → 检测 /[^0-9,\\s\\-]/ 判断字符输入\n"
                               "2. 图算法旧格式与新边列表格式需兼容 → 统一采用 [N, from, to, w, ...] 格式\n"
                               "3. 中缀转后缀随机数据生成垃圾字符 → 改为随机数字+运算符交替拼接",
        "AI使用记录": "使用 Claude Code 生成 55 条随机数据规则的分支逻辑\n"
                     "使用 GitHub Copilot 辅助正则表达式编写",
    },
    "2026.06.21": {
        "工作计划": "1. GraphVisualizer 重构（N 节点环形布局）\n"
                     "2. 更新 Dijkstra/BFS step 生成器适配边列表格式\n"
                     "3. 更新图算法 defaultData",
        "完成情况": "1. GraphVisualizer 重构：支持 N 节点环形布局\n"
                     "   · circularLayout(n) 函数：N 个节点均匀分布在圆周上\n"
                     "   · parseEdges(data) 函数：解析边列表格式 [N, from, to, w, ...]\n"
                     "   · 节点标签扩展到 A-P（16 个字母）\n"
                     "2. 更新 Dijkstra 和 BFS 的 step 生成器：解析边列表 → 构建邻接矩阵\n"
                     "3. 更新 GraphVisualizer 数据格式：data = [N, dist0, dist1, ...]\n"
                     "4. 更新图算法 defaultData：全部改为边列表格式（Dijkstra/BFS/DFS/Prim/Kruskal）\n"
                     "5. 更新随机数据生成器：图算法产生 4-7 个随机节点 + 随机边 + 随机权重\n"
                     "6. 更新 DataControl 图算法输入提示和数据显示标签",
        "出现的问题以及解决方案": "1. GraphVisualizer 边渲染 graphEdges 字段 weight 类型不一致 → 统一使用 GraphEdgeHighlight 的 weight 可选字段，e.weight ?? 1 兜底\n"
                               "2. Dijkstra meta 出现重复 description 属性 → 删除旧的 description，保留新的边列表格式说明\n"
                               "3. Prim 和 Kruskal 的 defaultData 复用同一句代码 → 分别定位替换",
        "AI使用记录": "使用 Claude Code 重构 GraphVisualizer，生成环形布局算法\n"
                     "AI 生成代码约 200 行，人工修改约 15%",
    },
    "2026.06.22": {
        "工作计划": "1. Electron 打包配置完善\n"
                     "2. 成功打包 Mac ARM64\n"
                     "3. 编写 README.md",
        "完成情况": "1. Electron 打包配置完善：\n"
                     "   · 修复 electron 版本号从 ^30.5.1 → 30.5.1\n"
                     "   · 修复 electron-builder 版本号从 ^24.13.3 → 24.13.3\n"
                     "   · 修复 tsconfig.node.json 移除 composite\n"
                     "   · 添加 Windows target 配置（nsis）\n"
                     "2. 成功打包 Mac ARM64：AlgoVisual-1.0.0-mac-arm64.dmg（92MB）+ .zip（89MB）\n"
                     "3. 整理源代码包：AlgoVisual-source.zip（110KB）\n"
                     "4. 编写 README.md：快速开始、项目结构、技术栈、功能说明",
        "出现的问题以及解决方案": "1. Windows 跨平台打包 electron-v30.5.1-win32-x64.zip 下载超时 → 建议在 Windows 环境构建\n"
                               "2. electron-builder 报 version is a range 错误 → 改为固定版本号\n"
                               "3. npm run build 中 tsc 命令找不到 → 重新 npm install",
        "AI使用记录": "使用 Claude Code 解决 electron-builder 配置问题\n"
                     "查阅 electron-builder 文档确认 ARM64 打包参数",
    },
    "2026.06.23": {
        "工作计划": "1. 最终版本整理与验证\n"
                     "2. 修复 package.json 版本号后重新打包\n"
                     "3. 参与最终项目文档和开发日志审核",
        "完成情况": "1. 确认 tsc --noEmit 零错误\n"
                     "2. 确认 vite build 通过（398KB JS + 18KB CSS）\n"
                     "3. 确认 55 个算法全部注册\n"
                     "4. 修复 package.json 版本号问题后重新打包\n"
                     "5. 更新源代码 zip 包（含全部文档）\n"
                     "6. 参与最终项目文档和开发日志的审核",
        "出现的问题以及解决方案": "最终阶段无明显新问题，主要为版本整理和文档审核工作。",
        "AI使用记录": "使用 Claude Code 整理文档格式",
    },
}


def add_styled_paragraph(doc, text, style_name, bold=False, size=None, font_name=None):
    """Add a paragraph with specified style and optional formatting overrides."""
    p = doc.add_paragraph(style=style_name)
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if size:
            run.font.size = size
        if font_name:
            run.font.name = font_name
    return p


def add_date_separator(doc, date_str):
    """Add a date separator like ☆☆☆☆☆2026.06.18☆☆☆☆☆"""
    p = doc.add_paragraph(style='Normal Indent')
    # Clear default run
    p.clear()

    # Stars before
    r1 = p.add_run('☆☆☆☆☆☆☆☆☆☆☆☆☆☆')
    r1.font.name = 'Weibei TC Bold'

    # Date
    r2 = p.add_run(date_str)
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.name = '宋体'
    r2.font.color.rgb = RGBColor(0, 0, 0)

    # Stars after
    r3 = p.add_run('☆☆☆☆☆☆☆☆☆☆☆☆☆☆☆☆☆')
    r3.font.name = 'Weibei TC Bold'

    return p


def add_day_section(doc, date_str, data):
    """Add a complete day log section."""
    add_date_separator(doc, date_str)

    for title in ["工作计划", "完成情况", "出现的问题以及解决方案", "AI使用记录"]:
        # Section header
        p = doc.add_paragraph()
        run = p.add_run(f"{title}：")
        run.bold = True
        run.font.size = Pt(12)

        # Content
        content = data.get(title, "（无记录）")
        doc.add_paragraph(content)

        # Empty line
        doc.add_paragraph()


def setup_page(doc):
    """Configure page layout."""
    for section in doc.sections:
        section.page_width  = Cm(21.0)  # A4
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)
        section.top_margin    = Cm(2.4)
        section.bottom_margin = Cm(2.4)


def main():
    # Load template to inherit styles
    template = Document(TEMPLATE)

    # Create new document from template (inherits styles)
    doc = Document(TEMPLATE)

    # Clear all content from template
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # Also remove empty paragraphs (some may remain)
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:p'):
            body.remove(child)

    setup_page(doc)

    # ── Title ──
    doc.add_heading('个人开发日志', level=1)

    # ── Name heading ──
    doc.add_heading('聂琦钢的开发日志', level=2)

    # ── Daily log entries ──
    for date_str in ["2026.06.18", "2026.06.19", "2026.06.20", "2026.06.21", "2026.06.22", "2026.06.23"]:
        data = DAILY_DATA[date_str]
        add_day_section(doc, date_str, data)

    # ── Separator ──
    doc.add_paragraph()

    # ── Project stats ──
    doc.add_heading('项目数据统计', level=2)
    doc.add_paragraph(
        "项目名称：算法过程可视化系统（AlgoVisual）\n"
        "担任角色：组长 / 项目架构 & 核心框架负责人\n"
        "负责模块：Electron 桌面端脚手架、TypeScript 类型系统、Zustand 全局状态机、项目配置与打包\n"
        "总任务数：7 项（T01-T07）\n"
        "开发周期：2026.06.18 - 2026.06.23（6天）\n"
        "AI 对话总轮次：约 20 轮\n"
        "AI 生成代码总行数：约 1200 行\n"
        "总人工修改比例：约 20%"
    )
    doc.add_paragraph()

    # ── Problems summary ──
    doc.add_heading('遇到的主要问题汇总', level=2)
    problems = [
        ("1. CJS/ESM 模块类型冲突", "postcss/tailwind 配置文件导出方式与 package.json type 字段不匹配", "配置文件改为 .mjs 扩展名", "✅ 已解决"),
        ("2. Electron 二进制下载失败", "GitHub Releases 在国内网络不稳定", "使用 ELECTRON_MIRROR 镜像", "✅ 已解决"),
        ("3. TypeScript 类型检查报 20+ 错误", "ActionType/VisualizerType 扩展后未同步更新所有引用", "逐一补全 ACTION_LABELS/ACTION_COLORS", "✅ 已解决"),
        ("4. 随机数据对所有算法生成相同格式", "未区分算法类型的输入需求", "按算法 ID 定制 55 条生成规则", "✅ 已解决"),
        ("5. electron-builder 版本范围报错", "electron-builder 需要固定版本号", "^30.5.1 → 30.5.1", "✅ 已解决"),
        ("6. Windows 跨平台打包网络超时", "下载 Windows Electron 二进制超时", "建议 Windows 环境构建", "⚠ 待解决"),
        ("7. node_modules 清理后 tsc 找不到", "npm run build 依赖 npx 调用全局 tsc", "重新 npm install", "✅ 已解决"),
    ]
    for title, desc, solution, status in problems:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(f"  问题描述：{desc}\n  解决方案：{solution}\n  状态：{status}")
        doc.add_paragraph()

    # ── AI usage summary ──
    doc.add_heading('AI 使用记录汇总', level=2)
    ai_records = [
        "日期       | AI 工具         | 使用场景             | 产出                              | 人工修改比例",
        "6.18      | Claude Code     | 项目脚手架搭建       | package.json/配置文件/index.html  | 20%",
        "6.18      | Claude Code     | 类型系统设计         | types/algo.ts（第一版）           | 15%",
        "6.18      | Claude Code     | Zustand 状态机       | AlgoContext.tsx                   | 20%",
        "6.18      | Claude Code     | Electron 主进程      | electron/main.ts                  | 10%",
        "6.19      | Claude Code     | 类型系统扩展         | ActionType 14种/VisualizerType 8种| 15%",
        "6.20      | Claude Code     | 随机数据生成器       | 55条分支规则                      | 25%",
        "6.20      | Claude Code     | 字符直输             | DataControl 符号检测逻辑          | 20%",
        "6.21      | Claude Code     | GraphVisualizer 重构 | 环形布局+N节点适配                | 15%",
        "6.22      | Claude Code     | 打包配置修复         | electron-builder 版本号修复       | 10%",
        "6.22      | GitHub Copilot  | TypeScript 补全      | 类型定义/接口实现代码片段         | 30%",
        "6.23      | Claude Code     | 文档整理             | 最终项目文档/开发日志             | 15%",
    ]
    for record in ai_records:
        doc.add_paragraph(record)
    doc.add_paragraph()
    doc.add_paragraph("AI 对话总轮次：约 20 轮 | AI 生成代码总行数：约 1200 行 | 总人工修改比例：约 20%")
    doc.add_paragraph()

    # ── Personal summary ──
    doc.add_heading('个人总结', level=2)
    summary = (
        "通过本项目，我完整实践了从前端工程化搭建到桌面应用打包的全流程。以下是我的核心收获：\n\n"
        "【技术方面】\n"
        "1. Electron + Vite + React + TypeScript 的技术栈组合在开发效率和运行性能之间取得了很好的平衡\n"
        "2. Zustand 的轻量级状态管理非常适合中等复杂度的桌面应用，特别是 getState() 在 setInterval 中的使用解决了 React 状态闭包问题\n"
        "3. TypeScript 严格模式下的接口先行设计有效避免了大量运行时错误，类型系统成为团队协作的「合约」\n"
        "4. electron-builder 的跨平台打包需要特别注意版本号的精确匹配和网络环境\n\n"
        "【工程方面】\n"
        "1. 好的项目脚手架是团队效率的基础——统一的 tsconfig、eslint、路径别名能避免大量配置碎片化问题\n"
        "2. 随机数据生成器看似简单，实际需要深入理解每个算法的输入约束才能生成合法数据\n"
        "3. 字符直输功能看似很小的改进，但需要同时修改 DataControl、AlgoContext 和多个算法的步骤生成器\n\n"
        "【协作方面】\n"
        "1. 作为架构负责人，需要提前定义好接口规范（VisualStep/AlgoStore 等）并确保团队成员理解一致\n"
        "2. 及时同步类型变更（如 ActionType 扩展）避免成员 B 和 C 的代码出现类型错误\n"
        "3. AI 工具在代码生成和文档编写方面效率极高，但架构决策和边界条件需要人工判断\n\n"
        "【不足与改进】\n"
        "1. 图算法跨平台打包未完成，后续可在 CI/CD 环境中配置多平台构建\n"
        "2. 可考虑引入离线 Electron 二进制包，避免网络依赖\n"
        "3. 代码签名证书的配置可以进一步提升用户体验"
    )
    doc.add_paragraph(summary)
    doc.add_paragraph()

    # ── PDF page images ──
    doc.add_heading('PDF 原件页面截图', level=2)

    if os.path.exists(IMG_DIR):
        images = sorted(
            [f for f in os.listdir(IMG_DIR) if f.endswith('.png')],
            key=lambda x: int(re.search(r'page-(\d+)', x).group(1))
        )

        for img_file in images:
            img_path = os.path.join(IMG_DIR, img_file)
            page_num = re.search(r'page-(\d+)', img_file).group(1)

            p = doc.add_paragraph()
            run = p.add_run(f'第 {page_num} 页')
            run.bold = True
            run.font.size = Pt(13)

            try:
                doc.add_picture(img_path, width=Inches(5.5))
                # Center align the last paragraph (image)
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f'[图片加载失败: {e}]')

            doc.add_paragraph()

    # ── Save ──
    doc.save(OUTPUT)
    print(f"Saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
